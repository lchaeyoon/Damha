import streamlit as st
import time
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_LINE_SPACING
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import os
import requests
import uuid
import json
import base64
from PIL import Image
from datetime import datetime
import io

# 페이지 설정
st.set_page_config(
    page_title="이미지 텍스트 추출 및 검수 시스템",
    page_icon="🔍",
    layout="wide"
)

# CSS 스타일
st.markdown("""
    <style>
    .stButton>button {
        width: 100%;
        background-color: #03a9f4;
        color: white;
    }
    .main {
        padding: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)

def clean_text(text):
    """텍스트 정리"""
    remove_chars = '☑◆●■□△▲▽▼→←↑↓★☆○◎◇◆□■△▲▽▼※~$'
    for char in remove_chars:
        text = text.replace(char, '')
    return ' '.join(text.split()).strip()

def extract_text_with_clova(image_bytes):
    """CLOVA OCR API를 사용한 텍스트 추출"""
    api_url = st.secrets["clova_ocr"]["api_url"]
    secret_key = st.secrets["clova_ocr"]["secret_key"]

    try:
        file_data_base64 = base64.b64encode(image_bytes).decode()
        
        request_json = {
            'images': [
                {
                    'format': 'jpg',
                    'name': 'demo',
                    'data': file_data_base64
                }
            ],
            'requestId': str(uuid.uuid4()),
            'version': 'V2',
            'timestamp': int(round(time.time() * 1000))
        }
            
        headers = {
            'X-OCR-SECRET': secret_key,
            'Content-Type': 'application/json'
        }
        
        response = requests.post(api_url, headers=headers, json=request_json)
        
        if response.status_code == 200:
            result = response.json()
            
            current_line = []
            lines = []
            last_y = None
            y_threshold = 10
            
            for image in result.get('images', []):
                fields = sorted(image.get('fields', []), 
                             key=lambda x: (x['boundingPoly']['vertices'][0]['y'], 
                                         x['boundingPoly']['vertices'][0]['x']))
                
                for field in fields:
                    if 'inferText' not in field:
                        continue
                        
                    text = clean_text(field['inferText'])
                    if not text:
                        continue
                        
                    current_y = field['boundingPoly']['vertices'][0]['y']
                    
                    if last_y is not None and abs(current_y - last_y) > y_threshold:
                        if current_line:
                            cleaned_line = clean_text(' '.join(current_line))
                            if cleaned_line:
                                lines.append(cleaned_line)
                            current_line = []
                    
                    current_line.append(text)
                    last_y = current_y
                
                if current_line:
                    cleaned_line = clean_text(' '.join(current_line))
                    if cleaned_line:
                        lines.append(cleaned_line)
            
            return '\n'.join(lines)
            
        else:
            st.error(f"API 오류: {response.status_code}")
            st.error(f"오류 메시지: {response.text}")
            return None
                
    except Exception as e:
        st.error(f"오류 발생: {str(e)}")
        return None

@st.cache_resource
def get_keywords_from_sheet():
    """구글 시트에서 키워드와 사유 가져오기"""
    try:
        scope = ['https://spreadsheets.google.com/feeds',
                'https://www.googleapis.com/auth/drive']
        
        credentials = {
            "type": st.secrets["gcp_service_account"]["type"],
            "project_id": st.secrets["gcp_service_account"]["project_id"],
            "private_key_id": st.secrets["gcp_service_account"]["private_key_id"],
            "private_key": st.secrets["gcp_service_account"]["private_key"],
            "client_email": st.secrets["gcp_service_account"]["client_email"],
            "client_id": st.secrets["gcp_service_account"]["client_id"],
            "auth_uri": st.secrets["gcp_service_account"]["auth_uri"],
            "token_uri": st.secrets["gcp_service_account"]["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["gcp_service_account"]["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["gcp_service_account"]["client_x509_cert_url"]
        }
        
        creds = ServiceAccountCredentials.from_json_keyfile_dict(credentials, scope)
        client = gspread.authorize(creds)
        
        sheet = client.open_by_url(st.secrets["spreadsheet"]["url"]).worksheet('키워드')
        
        keywords = sheet.col_values(2)[2:]
        reasons = sheet.col_values(3)[2:]
        
        keyword_notes = {}
        for keyword, reason in zip(keywords, reasons):
            if keyword.strip():
                keyword_notes[keyword] = reason if reason else ''
                
        return keyword_notes
        
    except Exception as e:
        st.error(f"구글 시트 데이터 가져오기 실패: {str(e)}")
        return None

def create_review_document(text, keyword_notes):
    """검수 결과 문서 생성"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = '맑은 고딕'
    
    lines = text.split('\n')
    
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        
        remaining_text = line
        current_pos = 0
        
        while remaining_text:
            earliest_keyword = None
            earliest_pos = len(remaining_text)
            
            for keyword in keyword_notes:
                pos = remaining_text.find(keyword)
                if pos != -1 and pos < earliest_pos:
                    earliest_keyword = keyword
                    earliest_pos = pos
            
            if earliest_keyword:
                if earliest_pos > 0:
                    run = paragraph.add_run(remaining_text[:earliest_pos])
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(10)
                
                run = paragraph.add_run(earliest_keyword)
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                
                if keyword_notes[earliest_keyword]:
                    run = paragraph.add_run(f" ({keyword_notes[earliest_keyword]}) ")
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 128, 0)
                
                remaining_text = remaining_text[earliest_pos + len(earliest_keyword):]
            else:
                if remaining_text:
                    run = paragraph.add_run(remaining_text)
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(10)
                break
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def main():
    st.title('🔍 이미지 텍스트 추출 및 검수 시스템')
    st.markdown('---')

    # 이미지 업로드
    uploaded_files = st.file_uploader(
        "이미지 파일을 업로드하세요 (여러 파일 선택 가능)",
        type=['png', 'jpg', 'jpeg'],
        accept_multiple_files=True
    )

    if uploaded_files:
        keyword_notes = get_keywords_from_sheet()
        if not keyword_notes:
            st.error("키워드 데이터를 가져올 수 없습니다.")
            st.info("구글 시트 연결을 확인해주세요.")
            return

        # 진행 상태 표시
        progress_text = "전체 진행 상황"
        progress_bar = st.progress(0)
        total_files = len(uploaded_files)

        for idx, uploaded_file in enumerate(uploaded_files):
            st.subheader(f"파일 처리 중: {uploaded_file.name}")
            
            # 이미지 미리보기
            image = Image.open(uploaded_file)
            st.image(image, caption=uploaded_file.name, use_column_width=True)
            
            # OCR 처리
            with st.spinner('텍스트 추출 중...'):
                image_bytes = uploaded_file.getvalue()
                extracted_text = extract_text_with_clova(image_bytes)
                
                if extracted_text:
                    st.success("텍스트 추출 완료")
                    
                    # 추출된 텍스트 표시
                    with st.expander("추출된 텍스트 보기"):
                        st.text_area("", extracted_text, height=200)
                    
                    # 검수 결과 문서 생성
                    with st.spinner('검수 결과 생성 중...'):
                        doc_io = create_review_document(extracted_text, keyword_notes)
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            # 다운로드 버튼
                            st.download_button(
                                label="📥 검수 결과 다운로드 (DOCX)",
                                data=doc_io.getvalue(),
                                file_name=f'검수결과_{os.path.splitext(uploaded_file.name)[0]}.docx',
                                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            )
                else:
                    st.error("텍스트를 추출할 수 없습니다.")
            
            # 진행 상태 업데이트
            progress_bar.progress((idx + 1) / total_files)
            st.markdown('---')

        st.success(f"모든 파일 처리 완료! (총 {total_files}개)")

    # 사용 방법
    with st.expander("사용 방법"):
        st.markdown("""
        ### 시스템 사용 방법
        1. 검수할 이미지 파일을 업로드합니다. (여러 파일 선택 가능)
        2. 시스템이 자동으로 다음 작업을 수행합니다:
           - OCR을 통한 텍스트 추출
           - 추출된 텍스트에서 키워드 검사
           - 검수 결과 문서 생성
        3. 각 파일별로 다음 정보를 확인할 수 있습니다:
           - 원본 이미지 미리보기
           - 추출된 텍스트 내용
           - 검수 결과 문서 다운로드
        4. 검수 결과 문서에서:
           - 빨간색: 검수 대상 키워드
           - 초록색: 검수 사유
        """)

if __name__ == '__main__':
    main() 
